import os
import io
import time
import base64
import requests
from pathlib import Path
from typing import List, Optional, Dict, Any
from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel
from sqlalchemy.orm import Session
from uuid import UUID

from app.db import get_db
from app.models import User, File as FileModel, OCRResult
from app.deps import get_current_active_user
from app.core.config import settings
from app.schemas import OCRPageData, OCRResult as OCRResultSchema, OCRResultCreate, OCRResultUpdate
from app.services.ai_extraction import process_document_with_ai, detect_document_type
from app.services.translation import process_mixed_language_text, is_translation_available

# OCR imports
try:
    from google.cloud import vision
    import fitz  # PyMuPDF
    from PIL import Image
    from docx import Document  # python-docx for DOCX files
    VISION_AVAILABLE = True
    DOCX_AVAILABLE = True
except ImportError as e:
    VISION_AVAILABLE = False
    DOCX_AVAILABLE = False

router = APIRouter()


class VisionAPIKeyClient:
    """Custom Vision API client using API key for authentication"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://vision.googleapis.com/v1"
    
    def text_detection(self, image):
        """Perform text detection using Google Vision API REST endpoint"""
        url = f"{self.base_url}/images:annotate?key={self.api_key}"
        
        # Prepare the request payload
        payload = {
            "requests": [
                {
                    "image": {
                        "content": base64.b64encode(image.content).decode('utf-8')
                    },
                    "features": [
                        {
                            "type": "TEXT_DETECTION",
                            "maxResults": 10
                        }
                    ]
                }
            ]
        }
        
        headers = {
            'Content-Type': 'application/json',
        }
        
        print(f"Making request to Vision API: {url}")
        print(f"Payload size: {len(str(payload))}")
        
        response = requests.post(url, json=payload, headers=headers, timeout=30)
        
        print(f"Vision API response status: {response.status_code}")
        if response.status_code != 200:
            print(f"Vision API error response: {response.text}")
        
        response.raise_for_status()
        
        result = response.json()
        print(f"Vision API response: {result}")
        
        return result


class OCRRequest(BaseModel):
    file_path: Optional[str] = None
    file_id: Optional[UUID] = None


class OCRPageResult(BaseModel):
    page: int
    text: str


class OCRResponse(BaseModel):
    pages: List[OCRPageResult]
    full_text: str
    total_pages: int
    file_path: str
    document_type: Optional[str] = None
    ai_extracted_data: Optional[Dict[str, Any]] = None
    translation_applied: Optional[bool] = None
    detected_languages: Optional[str] = None


def get_vision_client():
    """Initialize Google Vision client with proper credentials"""
    if not VISION_AVAILABLE:
        raise HTTPException(
            status_code=400,
            detail="Google Cloud Vision dependencies not installed"
        )
    
    # Try different credential approaches
    credentials_path = settings.GOOGLE_APPLICATION_CREDENTIALS
    
    # Approach 1: Service Account JSON file
    if credentials_path:
        # Handle both absolute and relative paths
        if not os.path.isabs(credentials_path):
            # Resolve relative to the backend root directory, not this file's location
            # __file__ is in: backend/app/api/api_v1/endpoints/ocr.py
            # We need to go up 5 levels to reach backend root
            backend_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))
            credentials_path = os.path.join(backend_root, credentials_path.lstrip('./'))
        
        credentials_path = os.path.abspath(credentials_path)
        if os.path.exists(credentials_path):
            os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = credentials_path
            try:
                return vision.ImageAnnotatorClient()
            except Exception as e:
                pass  # Fall through to API key approach
    
    # Approach 2: Try API Key authentication
    api_key = getattr(settings, 'GOOGLE_CLOUD_VISION_API_KEY', None)
    if api_key:
        try:
            # Use HTTP REST API directly with API key
            return VisionAPIKeyClient(api_key)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Failed to initialize Google Vision client with API key: {str(e)}. Please check your GOOGLE_CLOUD_VISION_API_KEY in the .env file."
            )
    
    # Approach 3: Try default credentials
    try:
        return vision.ImageAnnotatorClient()
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Google Cloud credentials not configured. Please provide either:\n1. GOOGLE_APPLICATION_CREDENTIALS path to JSON file\n2. GOOGLE_CLOUD_VISION_API_KEY in .env\n3. Default application credentials\nError: {str(e)}"
        )


def extract_text_from_image(image_data: bytes, client, is_pdf_page: bool = False) -> str:
    """Extract text from image using Google Vision API with enhanced detection"""
    try:
        # Handle both regular Vision client and our custom API key client
        if isinstance(client, VisionAPIKeyClient):
            # Use our custom API key client - pass raw image data
            image_obj = type('Image', (), {'content': image_data})()
            response_json = client.text_detection(image_obj)
            
            # Parse the REST API response
            if 'responses' in response_json and response_json['responses']:
                response = response_json['responses'][0]
                if 'error' in response:
                    raise Exception(f"Vision API error: {response['error']['message']}")
                
                if 'textAnnotations' in response and response['textAnnotations']:
                    return response['textAnnotations'][0]['description']
            return ""
        else:
            # Use regular Vision client
            image = vision.Image(content=image_data)
            
            # Use document_text_detection for better layout preservation, especially for PDFs
            if is_pdf_page:
                response = client.document_text_detection(image=image)
                
                if response.error.message:
                    raise Exception(f"Vision API error: {response.error.message}")
                
                # For document text detection, use full_text_annotation for better layout preservation
                if response.full_text_annotation:
                    return response.full_text_annotation.text
                # Fallback to text_annotations if full_text_annotation is not available
                elif response.text_annotations:
                    return response.text_annotations[0].description
                return ""
            else:
                # For regular images, use standard text detection
                response = client.text_detection(image=image)
                
                if response.error.message:
                    raise Exception(f"Vision API error: {response.error.message}")
                
                # Get full text annotation
                texts = response.text_annotations
                if texts:
                    return texts[0].description
                return ""
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text from image: {str(e)}"
        )


def pdf_to_images_and_extract_text(pdf_path: str, client: vision.ImageAnnotatorClient) -> List[OCRPageResult]:
    """Convert PDF pages to images and extract text from each page"""
    try:
        doc = fitz.open(pdf_path)
        pages_text = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # Convert page to image with reasonable DPI (150 is good balance of quality/speed)
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom = ~150 DPI
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PIL Image then to bytes
            img_data = pix.tobytes("png")
            
            # Extract text using Vision API with enhanced detection for PDFs
            text = extract_text_from_image(img_data, client, is_pdf_page=True)
            
            pages_text.append(OCRPageResult(page=page_num + 1, text=text))
            
        doc.close()
        return pages_text
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process PDF: {str(e)}"
        )


def extract_text_from_docx(docx_path: str) -> List[OCRPageResult]:
    """Extract text from DOCX file (no OCR needed)"""
    try:
        if not DOCX_AVAILABLE:
            raise HTTPException(
                status_code=400,
                detail="DOCX processing not available. python-docx not installed."
            )
        
        doc = Document(docx_path)
        text_content = []
        
        # Extract all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        full_text = "\n".join(text_content)
        return [OCRPageResult(page=1, text=full_text)]
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process DOCX: {str(e)}"
        )


def process_image_file(image_path: str, client: vision.ImageAnnotatorClient) -> List[OCRPageResult]:
    """Process image file (JPG, PNG, TIFF) using Vision API"""
    try:
        # Handle TIFF files - convert to PNG if needed
        file_ext = Path(image_path).suffix.lower()
        if file_ext in ['.tiff', '.tif']:
            # Convert TIFF to PNG for Vision API
            with Image.open(image_path) as img:
                # Convert to RGB if needed
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Convert to bytes
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='PNG')
                image_data = img_bytes.getvalue()
        else:
            # For JPG, PNG - read directly
            with open(image_path, 'rb') as f:
                image_data = f.read()
        
        text = extract_text_from_image(image_data, client, is_pdf_page=False)
        return [OCRPageResult(page=1, text=text)]
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process image file: {str(e)}"
        )


def validate_file_access(file_path: str, current_user: User) -> str:
    """Validate that the file exists and user has access to it"""
    if not file_path:
        raise HTTPException(status_code=400, detail="File path is required")
    
    # Convert to absolute path
    abs_path = os.path.abspath(file_path)
    storage_dir = os.path.abspath(settings.STORAGE_DIR)
    
    # Ensure file is within storage directory (security check)
    if not abs_path.startswith(storage_dir):
        raise HTTPException(status_code=403, detail="Access denied to file outside storage directory")
    
    # Check if file exists
    if not os.path.exists(abs_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Validate file type
    file_ext = Path(abs_path).suffix.lower()
    if file_ext not in ['.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.tif', '.docx', '.doc']:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Supported types: PDF, JPG, JPEG, PNG, TIFF, DOCX, DOC"
        )
    
    return abs_path


@router.post("/extract_text", response_model=OCRResponse)
async def extract_text(
    request: OCRRequest,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Extract text from PDF or image files using Google Vision API"""
    start_time = time.time()
    
    # Validate input
    if not request.file_path and not request.file_id:
        raise HTTPException(
            status_code=400,
            detail="Either file_path or file_id must be provided"
        )
    
    # Handle file_id lookup
    file_record = None
    if request.file_id:
        file_record = db.query(FileModel).filter(
            FileModel.id == request.file_id,
            FileModel.uploaded_by == current_user.id
        ).first()
        if not file_record:
            raise HTTPException(status_code=404, detail="File not found")
        # Always validate file access even when using file_id
        abs_file_path = validate_file_access(file_record.file_path, current_user)
    else:
        # Validate file access using file_path
        abs_file_path = validate_file_access(request.file_path, current_user)
        # Try to find existing file record by path
        file_record = db.query(FileModel).filter(
            FileModel.file_path == abs_file_path,
            FileModel.uploaded_by == current_user.id
        ).first()
    
    # Check if OCR already exists for this file
    if file_record:
        existing_ocr = db.query(OCRResult).filter(
            OCRResult.file_id == file_record.id
        ).first()
        if existing_ocr:
            # Return existing OCR result
            pages_data = [OCRPageResult(page=p["page"], text=p["text"]) for p in existing_ocr.blocks_json or []]
            return OCRResponse(
                pages=pages_data,
                full_text=existing_ocr.edited_text or existing_ocr.raw_text or "",
                total_pages=len(pages_data),
                file_path=request.file_path or file_record.file_path
            )
    
    # Determine file type and process accordingly
    try:
        file_ext = Path(abs_file_path).suffix.lower()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file path: {str(e)}")
    
    try:
        if file_ext == '.pdf':
            # Initialize Vision client for PDF processing
            vision_client = get_vision_client()
            pages_text = pdf_to_images_and_extract_text(abs_file_path, vision_client)
            
        elif file_ext in ['.docx', '.doc']:
            # Process DOCX/DOC files (direct text extraction, no OCR needed)
            pages_text = extract_text_from_docx(abs_file_path)
            
        else:
            # Process image files (JPG, PNG, TIFF)
            vision_client = get_vision_client()
            pages_text = process_image_file(abs_file_path, vision_client)
        
        # Combine all text - handle potential Unicode issues
        try:
            raw_full_text = "\n\n".join([page.text for page in pages_text if page.text.strip()])
        except Exception as e:
            raw_full_text = ""
        
        # Process with translation if available
        translation_applied = False
        detected_languages = "unknown"
        processed_full_text = raw_full_text
        
        if is_translation_available() and raw_full_text.strip():
            try:
                processed_full_text, detected_languages = process_mixed_language_text(raw_full_text)
                translation_applied = detected_languages != "en" and "si" in detected_languages
            except Exception as e:
                pass  # Continue without translation
        
        # Process with AI extraction
        ai_extracted_data = None
        document_type = None
        
        if processed_full_text.strip():
            try:
                ai_result = process_document_with_ai(processed_full_text)
                document_type = ai_result.get("document_type")
                ai_extracted_data = {
                    "document_type": ai_result.get("document_type"),
                    "extracted_data": ai_result.get("extracted_data", {}),
                    "general_data": ai_result.get("general_data", {})
                }
            except Exception as e:
                ai_extracted_data = {"error": f"AI processing failed: {str(e)}"}
        
        # Calculate processing time
        processing_time = int((time.time() - start_time) * 1000)
        
        # Save OCR result to database if we have a file record
        if file_record:
            try:
                # Convert pages_text to JSON format for database
                pages_data_json = [{"page": page.page, "text": page.text} for page in pages_text]
                
                ocr_result = OCRResult(
                    raw_text=processed_full_text,  # Store processed (translated) text
                    blocks_json=pages_data_json,  # Use blocks_json instead of pages_data
                    confidence_score=85,  # Placeholder - Google Vision doesn't provide overall confidence
                    processing_time=processing_time,
                    ocr_engine="google_vision",
                    language=detected_languages,  # Use language instead of language_detected
                    file_id=file_record.id,
                    processed_by=current_user.id
                )
                db.add(ocr_result)
                db.commit()
                db.refresh(ocr_result)
            except Exception as e:
                # Don't fail the whole request if database save fails
                pass
        
        return OCRResponse(
            pages=pages_text,
            full_text=processed_full_text,
            total_pages=len(pages_text),
            file_path=request.file_path or (file_record.file_path if file_record else abs_file_path),
            document_type=document_type,
            ai_extracted_data=ai_extracted_data,
            translation_applied=translation_applied,
            detected_languages=detected_languages
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"OCR processing failed: {str(e)}"
        )


@router.get("/results/{file_id}", response_model=OCRResultSchema)
def get_ocr_result(
    file_id: UUID,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Get OCR result for a specific file"""
    # Verify file belongs to user
    file_record = db.query(FileModel).filter(
        FileModel.id == file_id,
        FileModel.uploaded_by == current_user.id
    ).first()
    if not file_record:
        raise HTTPException(status_code=404, detail="File not found")
    
    # Get OCR result
    ocr_result = db.query(OCRResult).filter(
        OCRResult.file_id == file_id
    ).first()
    if not ocr_result:
        raise HTTPException(status_code=404, detail="OCR result not found")
    
    return ocr_result


@router.put("/results/{ocr_id}", response_model=OCRResultSchema)
def update_ocr_result(
    ocr_id: UUID,
    update_data: OCRResultUpdate,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Update OCR result (for user editing)"""
    # Get OCR result and verify user access
    ocr_result = db.query(OCRResult).join(FileModel).filter(
        OCRResult.id == ocr_id,
        FileModel.uploaded_by == current_user.id
    ).first()
    if not ocr_result:
        raise HTTPException(status_code=404, detail="OCR result not found")
    
    # Update fields
    update_dict = update_data.model_dump(exclude_unset=True)
    for field, value in update_dict.items():
        setattr(ocr_result, field, value)
    
    # Mark as edited if text was changed
    if update_data.edited_text is not None:
        ocr_result.is_edited = True
    
    db.commit()
    db.refresh(ocr_result)
    return ocr_result


@router.post("/analyze_document", response_model=Dict[str, Any])
async def analyze_document(
    request: OCRRequest,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Analyze document and extract structured data using AI (without running OCR again)"""
    
    # Get OCR result from database first
    file_record = None
    if request.file_id:
        file_record = db.query(FileModel).filter(
            FileModel.id == request.file_id,
            FileModel.uploaded_by == current_user.id
        ).first()
        if not file_record:
            raise HTTPException(status_code=404, detail="File not found")
    
    # Get existing OCR result
    if file_record:
        ocr_result = db.query(OCRResult).filter(
            OCRResult.file_id == file_record.id
        ).first()
        if not ocr_result:
            raise HTTPException(status_code=404, detail="OCR result not found. Please run OCR first.")
        
        ocr_text = ocr_result.edited_text or ocr_result.raw_text
    else:
        raise HTTPException(status_code=400, detail="file_id is required for document analysis")
    
    if not ocr_text:
        raise HTTPException(status_code=400, detail="No text available for analysis")
    
    try:
        # Process with AI extraction
        ai_result = process_document_with_ai(ocr_text)
        
        return {
            "file_id": str(file_record.id),
            "document_analysis": ai_result,
            "ocr_text_length": len(ocr_text),
            "analysis_timestamp": time.time()
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Document analysis failed: {str(e)}"
        )


@router.post("/extract_utilities")
async def extract_utilities_from_document(
    request: OCRRequest,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Extract utilities information from OCR text using AI"""
    
    # Get OCR text from file
    ocr_text = None
    file_record = None
    
    if request.file_id:
        try:
            file_record = db.query(FileModel).filter(
                FileModel.id == UUID(str(request.file_id)),
                FileModel.uploaded_by == current_user.id
            ).first()
        except Exception as e:
            raise HTTPException(status_code=400, detail="Invalid file_id format")
    
    # Get existing OCR result
    if file_record:
        ocr_result = db.query(OCRResult).filter(
            OCRResult.file_id == file_record.id
        ).first()
        if not ocr_result:
            raise HTTPException(status_code=404, detail="OCR result not found. Please run OCR first.")
        
        ocr_text = ocr_result.edited_text or ocr_result.raw_text
    else:
        raise HTTPException(status_code=400, detail="file_id is required for utilities extraction")
    
    if not ocr_text:
        raise HTTPException(status_code=400, detail="No text available for utilities extraction")
    
    try:
        # Import the utilities extraction function
        from app.services.ai_extraction import extract_utilities_data
        
        # Extract utilities data
        utilities_result = extract_utilities_data(ocr_text)
        
        return {
            "file_id": str(file_record.id),
            "utilities_data": utilities_result,
            "ocr_text_length": len(ocr_text),
            "extraction_timestamp": time.time(),
            "extraction_type": "utilities"
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Utilities extraction failed: {str(e)}"
        )


class BatchOCRRequest(BaseModel):
    file_ids: List[UUID]
    document_type: Optional[str] = None
    
    
class BatchOCRResponse(BaseModel):
    files_processed: int
    total_files: int
    results: List[Dict[str, Any]]
    consolidated_data: Optional[Dict[str, Any]] = None
    processing_time_seconds: float


@router.post("/batch_extract_text", response_model=BatchOCRResponse)
async def batch_extract_text(
    request: BatchOCRRequest,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """
    Process multiple files with OCR and AI analysis simultaneously.
    This provides more accurate results for property documents that span multiple pages/files.
    """
    start_time = time.time()
    
    if not request.file_ids:
        raise HTTPException(status_code=400, detail="No file IDs provided")
    
    if len(request.file_ids) > 10:  # Limit batch size for performance
        raise HTTPException(status_code=400, detail="Maximum 10 files per batch")
    
    # Verify all files exist and belong to the user
    files = []
    for file_id in request.file_ids:
        file_record = db.query(FileModel).filter(
            FileModel.id == file_id
        ).first()
        
        if not file_record:
            raise HTTPException(status_code=404, detail=f"File {file_id} not found")
        
        files.append(file_record)
    
    # Initialize Vision client for batch processing
    vision_client = get_vision_client()
    
    # Process each file with OCR
    results = []
    all_extracted_text = []
    
    for file_record in files:
        try:
            # Validate file access
            abs_file_path = validate_file_access(file_record.file_path, current_user)
            file_ext = Path(abs_file_path).suffix.lower()
            
            # Process individual file based on type
            if file_ext == '.pdf':
                pages_text = pdf_to_images_and_extract_text(abs_file_path, vision_client)
                text = "\n\n".join([page.text for page in pages_text if page.text.strip()])
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif']:
                pages_text = process_image_file(abs_file_path, vision_client)
                text = pages_text[0].text if pages_text else ""
            elif file_ext in ['.docx', '.doc']:
                pages_text = extract_text_from_docx(abs_file_path)
                text = pages_text[0].text if pages_text else ""
            else:
                continue
                
            if text and text.strip():
                # Get document type and AI analysis for individual file
                try:
                    doc_type = detect_document_type(text)
                    ai_data = process_document_with_ai(text)
                except Exception as ai_error:
                    doc_type = "unknown"
                    ai_data = {"error": f"AI analysis failed: {str(ai_error)}"}
                
                result = {
                    "file_id": str(file_record.id),
                    "filename": file_record.filename,
                    "extracted_text": text,
                    "document_type": doc_type,
                    "ai_extracted_data": ai_data,
                    "text_length": len(text)
                }
                
                results.append(result)
                all_extracted_text.append(text)
                
        except Exception as e:
            # Add error info but continue processing other files
            results.append({
                "file_id": str(file_record.id),
                "filename": file_record.filename,
                "error": str(e),
                "success": False
            })
    
    # Consolidate data from all files for more accurate analysis
    consolidated_data = None
    if all_extracted_text and len(all_extracted_text) > 1:
        try:
            # Combine all text for comprehensive analysis
            combined_text = "\n\n--- DOCUMENT SEPARATOR ---\n\n".join(all_extracted_text)
            
            # Detect overall document type from combined text
            try:
                combined_doc_type = detect_document_type(combined_text)
            except Exception:
                combined_doc_type = "property_documents"
            
            # Process combined text with AI for better accuracy
            try:
                consolidated_ai_data = process_document_with_ai(combined_text)
            except Exception as e:
                consolidated_ai_data = {"error": f"AI processing failed: {str(e)}"}
            
            consolidated_data = {
                "document_type": combined_doc_type,
                "ai_extracted_data": consolidated_ai_data,
                "total_text_length": len(combined_text),
                "files_analyzed": len(all_extracted_text)
            }
            
        except Exception as e:
            consolidated_data = {"error": f"Consolidation failed: {str(e)}"}
    
    processing_time = time.time() - start_time
    
    return BatchOCRResponse(
        files_processed=len([r for r in results if "error" not in r]),
        total_files=len(files),
        results=results,
        consolidated_data=consolidated_data,
        processing_time_seconds=processing_time
    )