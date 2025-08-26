"""
Professional Template Engine for Valuation Reports

Maps database entities to template placeholders and generates professional
DOCX and PDF reports following Sri Lankan valuation standards.
"""

import os
import re
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn

from sqlalchemy.orm import Session
from app.models import (
    Report, User, ValuerProfile, Client, Property, 
    Identification, Location, Access, Site, Building,
    Utilities, Planning, Locality, ValuationLine,
    ValuationSummary, Disclaimer, Certificate, Appendix
)
from app.services.number_to_words import convert_lkr_to_currency_words


class TemplateDataMapper:
    """Maps database entities to template placeholder variables"""
    
    def __init__(self, db: Session):
        self.db = db
    
    def map_report_data(self, report_id: str) -> Dict[str, Any]:
        """
        Map all database entities for a report to template variables
        
        Returns a complete data dictionary matching the template structure
        from report structure.md
        """
        # Get main report with all relationships
        report = self.db.query(Report).filter(Report.id == report_id).first()
        if not report:
            raise ValueError(f"Report {report_id} not found")
        
        # Get related entities
        author = report.author
        client = report.client
        properties = self.db.query(Property).filter(Property.report_id == report_id).all()
        
        # Build complete template data
        template_data = {}
        
        # Map valuer profile
        template_data.update(self._map_valuer_profile(author))
        
        # Map report metadata
        template_data.update(self._map_report_metadata(report, client))
        
        # Map properties (assume single property for now, can extend for multi-property)
        if properties:
            property_data = self._map_property_data(properties[0])
            template_data.update(property_data)
        
        # Map valuation summary
        template_data.update(self._map_valuation_data(report))
        
        # Map executive summary (auto-generated)
        template_data.update(self._generate_executive_summary(template_data))
        
        return template_data
    
    def _map_valuer_profile(self, user: User) -> Dict[str, Any]:
        """Map valuer profile data to template variables"""
        profile = user.valuer_profile if user else None
        
        return {
            'valuer': {
                'full_name': user.full_name if user else '',
                'titles': profile.titles if profile and profile.titles else '',
                'qualifications': ', '.join(profile.qualifications) if profile and profile.qualifications else 'Chartered Valuer',
                'panels_list': ', '.join(profile.panels) if profile and profile.panels else 'Panel Valuer',
                'address_multiline': profile.address.replace(',', '\n') if profile and profile.address else '',
                'phones_list': ', '.join(profile.phones) if profile and profile.phones else '',
                'email': profile.email if profile and profile.email else user.email if user else '',
                'registration_no': profile.registration_no if profile and profile.registration_no else 'IVSL Member'
            }
        }
    
    def _map_report_metadata(self, report: Report, client: Client) -> Dict[str, Any]:
        """Map report metadata to template variables"""
        return {
            'report': {
                'ref': report.ref or f'VRP-{datetime.now().strftime("%Y%m%d")}-{str(report.id)[:8].upper()}',
                'date': report.report_date.strftime('%d %B %Y') if report.report_date else datetime.now().strftime('%d %B %Y'),
                'purpose': report.purpose or 'Banking and General Purposes'
            },
            'client': {
                'name': client.name if client else 'Client Name Not Specified'
            },
            'inspection': {
                'date': report.inspection_date.strftime('%d %B %Y') if report.inspection_date else 'To be confirmed',
                'by': report.author.full_name if report.author else 'Licensed Valuer'
            },
            'basis_of_value': report.basis_of_value or 'Market Value'
        }
    
    def _map_property_data(self, property_obj: Property) -> Dict[str, Any]:
        """Map property and related entity data to template variables"""
        # Get related entities
        identification = self.db.query(Identification).filter(Identification.property_id == property_obj.id).first()
        location = self.db.query(Location).filter(Location.property_id == property_obj.id).first()
        access = self.db.query(Access).filter(Access.property_id == property_obj.id).first()
        site = self.db.query(Site).filter(Site.property_id == property_obj.id).first()
        buildings = self.db.query(Building).filter(Building.property_id == property_obj.id).all()
        utilities = self.db.query(Utilities).filter(Utilities.property_id == property_obj.id).first()
        planning = self.db.query(Planning).filter(Planning.property_id == property_obj.id).first()
        locality = self.db.query(Locality).filter(Locality.property_id == property_obj.id).first()
        
        template_data = {}
        
        # Identification data
        if identification:
            template_data['id'] = {
                'lot_number': identification.lot_number or '',
                'plan_number': identification.plan_number or '',
                'plan_date': identification.plan_date.strftime('%d/%m/%Y') if identification.plan_date else '',
                'surveyor_name': identification.surveyor_name or '',
                'land_name': identification.land_name or '',
                'extent_local': identification.extent_local or f'{identification.extent_perches} Perches' if identification.extent_perches else '',
                'extent_metric': f'{identification.extent_sqm} sqm' if identification.extent_sqm else '',
                'extent_perches': identification.extent_perches or 0,
                'boundaries': identification.boundaries or {}
            }
            
            # Title information
            template_data['title'] = {
                'owner_name': identification.title_owner or 'Owner Name Not Specified',
                'deed_no': identification.deed_no or 'N/A',
                'deed_date': identification.deed_date.strftime('%d %B %Y') if identification.deed_date else 'N/A',
                'notary': identification.notary or 'N/A',
                'interest': identification.interest or 'Freehold'
            }
        
        # Location data
        if location:
            template_data['location'] = {
                'village': location.village or '',
                'district': location.district or '',
                'province': location.province or '',
                'gn_division': location.gn_division or '',
                'ds_division': location.ds_division or '',
                'lat': location.lat,
                'lng': location.lng
            }
            
            template_data['address'] = {
                'full': location.address_full or 'Address not specified'
            }
        
        # Access data
        if access:
            template_data['access'] = {
                'landmark': access.landmark or 'Local landmark',
                'directions': access.directions_text or 'Access directions available upon inspection',
                'road_names': access.road_names or 'Local roads',
                'road_width': access.road_width or 'Standard width',
                'road_surface': access.road_surface or 'Metaled',
                'maintainer': access.maintainer or 'Local Authority'
            }
        
        # Site data
        if site:
            template_data['site'] = {
                'shape': site.shape or 'Regular',
                'topography': site.topography or 'Level',
                'level_vs_road': site.level_vs_road or 'Level with road',
                'soil': site.soil or 'Good bearing soil',
                'water_table_depth': f'{site.water_table_depth_ft} ft' if site.water_table_depth_ft else 'Adequate depth',
                'frontage': f'{site.frontage_ft} ft' if site.frontage_ft else 'To be measured',
                'features_list': ', '.join(site.features) if site.features else 'Standard features',
                'flood_risk': site.flood_risk or 'Low risk'
            }
        
        # Buildings data
        template_data['buildings'] = []
        for i, building in enumerate(buildings):
            building_data = {
                'index': i + 1,
                'type': building.type or 'Residential building',
                'storeys': f'{building.storeys}' if building.storeys else 'Single',
                'structure': building.structure or 'Masonry',
                'roof': {
                    'type': building.roof_type or 'Tiled',
                    'structure': building.roof_structure or 'Timber'
                },
                'walls': building.walls or 'Masonry walls',
                'floors': building.floors or 'Cement floors',
                'doors': building.doors or 'Timber doors',
                'windows': building.windows or 'Timber windows',
                'layout': building.layout_text or 'Functional layout',
                'area_sqft': building.area_sqft or 0,
                'area_sqm': building.area_sqm or 0,
                'age': f'{building.age_years}' if building.age_years else 'Moderate',
                'condition': building.condition or 'Good',
                'occupancy': building.occupancy or 'Owner occupied'
            }
            template_data['buildings'].append(building_data)
        
        # Utilities data
        if utilities:
            template_data['utilities'] = {
                'electricity': 'Available' if utilities.electricity else 'Not available',
                'water': 'Available' if utilities.water else 'Not available',
                'telecom': 'Available' if utilities.telecom else 'Not available',
                'sewerage': utilities.sewerage or 'Standard',
                'drainage': utilities.drainage or 'Adequate',
                'other': utilities.other or 'Standard utilities'
            }
        
        # Planning data
        if planning:
            template_data['planning'] = {
                'zoning': planning.zoning or 'Residential',
                'street_line': planning.street_line or 'Standard setbacks',
                'easements': planning.easements or 'No adverse easements identified'
            }
        
        # Locality data
        if locality:
            template_data['locality'] = {
                'narrative': locality.narrative or 'Established residential area with good amenities and infrastructure.'
            }
        
        return template_data
    
    def _map_valuation_data(self, report: Report) -> Dict[str, Any]:
        """Map valuation data including calculations and summary"""
        # Get valuation lines
        properties = self.db.query(Property).filter(Property.report_id == report.id).all()
        valuation_data = {
            'valuation': {
                'approach': 'Comparison Method',
                'market_value_numeric': 0,
                'market_value_words': '',
                'forced_sale_value_numeric': 0,
                'fsv_pct': report.fsv_percentage or 80
            },
            'calc': {
                'land_rate_perch': 0,
                'land_value': 0,
                'other_value': 0
            },
            'evidence': {
                'low': 'Market research ongoing',
                'high': 'Market research ongoing',
                'size_range': 'Similar properties'
            }
        }
        
        if properties:
            # Get valuation summary
            valuation_summary = self.db.query(ValuationSummary).filter(
                ValuationSummary.report_id == report.id
            ).first()
            
            if valuation_summary:
                market_value = valuation_summary.market_value
                forced_sale_value = valuation_summary.forced_sale_value
                
                valuation_data['valuation'].update({
                    'market_value_numeric': market_value,
                    'market_value_words': convert_lkr_to_currency_words(market_value),
                    'forced_sale_value_numeric': forced_sale_value
                })
            
            # Get individual valuation lines
            for property_obj in properties:
                lines = self.db.query(ValuationLine).filter(
                    ValuationLine.property_id == property_obj.id
                ).order_by(ValuationLine.sort_order).all()
                
                for line in lines:
                    if line.line_type == 'land':
                        # Get identification for perch calculation
                        identification = self.db.query(Identification).filter(
                            Identification.property_id == property_obj.id
                        ).first()
                        
                        if identification and identification.extent_perches:
                            rate_per_perch = line.rate
                            valuation_data['calc'].update({
                                'land_rate_perch': rate_per_perch,
                                'land_value': line.value
                            })
        
        return valuation_data
    
    def _generate_executive_summary(self, template_data: Dict[str, Any]) -> Dict[str, Any]:
        """Auto-generate executive summary from other sections"""
        summary = {
            'summary': {
                'property_type': 'Residential Property',
                'buildings_brief': 'Property details as per inspection',
            }
        }
        
        # Extract property type from buildings
        if template_data.get('buildings') and len(template_data['buildings']) > 0:
            building = template_data['buildings'][0]
            summary['summary']['property_type'] = building.get('type', 'Residential Property')
            
            area = building.get('area_sqft', 0)
            if area > 0:
                summary['summary']['buildings_brief'] = f"{building.get('storeys', 'Single')}-storey {building.get('type', 'building')} of approximately {area} sq.ft"
        
        return summary


class ProfessionalTemplateEngine:
    """Professional template engine for generating valuation reports"""
    
    def __init__(self, db: Session):
        self.db = db
        self.mapper = TemplateDataMapper(db)
        self.templates_dir = Path(__file__).parent.parent / "templates"
        self.ensure_templates_dir()
    
    def ensure_templates_dir(self):
        """Ensure templates directory exists"""
        self.templates_dir.mkdir(exist_ok=True)
    
    def generate_professional_docx(self, report_id: str) -> BytesIO:
        """
        Generate a professional DOCX report following the report structure template
        
        Args:
            report_id: The report ID to generate
            
        Returns:
            BytesIO: DOCX content as bytes
        """
        # Get mapped template data
        template_data = self.mapper.map_report_data(report_id)
        
        # Create new document
        doc = Document()
        
        # Set up document properties
        self._setup_document_styles(doc)
        
        # Generate report sections following the template structure
        self._add_cover_page(doc, template_data)
        self._add_executive_summary(doc, template_data)
        self._add_introduction(doc, template_data)
        self._add_property_identification(doc, template_data)
        self._add_location_description(doc, template_data)
        self._add_access_description(doc, template_data)
        self._add_site_description(doc, template_data)
        self._add_improvements_description(doc, template_data)
        self._add_services_utilities(doc, template_data)
        self._add_planning_zoning(doc, template_data)
        self._add_locality_context(doc, template_data)
        self._add_valuation_approach(doc, template_data)
        self._add_valuation_calculations(doc, template_data)
        self._add_conclusion(doc, template_data)
        self._add_certificate_identity(doc, template_data)
        self._add_assumptions_conditions(doc, template_data)
        self._add_signature_block(doc, template_data)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _setup_document_styles(self, doc: Document):
        """Set up professional document styles"""
        # Configure styles for professional appearance
        styles = doc.styles
        
        # Title style
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', 1)
            title_font = title_style.font
            title_font.name = 'Times New Roman'
            title_font.size = Pt(16)
            title_font.bold = True
        
        # Heading style
        if 'CustomHeading' not in styles:
            heading_style = styles.add_style('CustomHeading', 1) 
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.size = Pt(14)
            heading_font.bold = True
    
    def _add_cover_page(self, doc: Document, data: Dict[str, Any]):
        """Add professional cover page / letterhead"""
        valuer = data.get('valuer', {})
        report = data.get('report', {})
        
        # Left header - Valuer information
        valuer_para = doc.add_paragraph()
        valuer_run = valuer_para.add_run(f"{valuer.get('titles', '')} {valuer.get('full_name', '')}\n")
        valuer_run.font.bold = True
        valuer_run.font.size = Pt(12)
        
        quals_run = valuer_para.add_run(f"{valuer.get('qualifications', '')}\n")
        panels_run = valuer_para.add_run(f"Panel Valuer: {valuer.get('panels_list', '')}")
        
        # Right header - Contact information
        doc.add_paragraph()
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        contact_text = valuer.get('address_multiline', '') + '\n'
        contact_text += f"Tel: {valuer.get('phones_list', '')}\n"
        contact_text += f"Email: {valuer.get('email', '')}"
        
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = Pt(10)
        
        # Reference and date line
        doc.add_paragraph()
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_text = f"My Ref: {report.get('ref', '')}    Date: {report.get('date', '')}"
        ref_run = ref_para.add_run(ref_text)
        ref_run.font.bold = True
        
        # Main title
        doc.add_paragraph()
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("VALUATION REPORT")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        
        # Subject line
        doc.add_paragraph()
        subject_para = doc.add_paragraph()
        subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        id_data = data.get('id', {})
        subject_text = f"Valuation Report of the Property depicted as Lot {id_data.get('lot_number', 'XX')} "
        subject_text += f"in Plan No. {id_data.get('plan_number', 'XXXX')} "
        subject_text += f"dated {id_data.get('plan_date', 'XX/XX/XXXX')}, "
        subject_text += f"prepared by {id_data.get('surveyor_name', 'Licensed Surveyor')}."
        
        subject_run = subject_para.add_run(subject_text)
        subject_run.font.size = Pt(12)
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, data: Dict[str, Any]):
        """Add executive summary table"""
        heading = doc.add_heading('Executive Summary', level=1)
        
        # Create summary table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # Summary data
        summary = data.get('summary', {})
        location = data.get('location', {})
        id_data = data.get('id', {})
        valuation = data.get('valuation', {})
        
        rows_data = [
            ['Property Type', summary.get('property_type', 'Residential Property')],
            ['Location', f"{location.get('village', '')}, {location.get('district', '')} District, {location.get('province', '')}"],
            ['Land Extent', f"{id_data.get('extent_local', '')} (≈ {id_data.get('extent_metric', '')})"],
            ['Open Market Value (OMV)', f"Rs. {valuation.get('market_value_numeric', 0):,}"],
            ['Forced Sale Value (FSV)', f"Rs. {valuation.get('forced_sale_value_numeric', 0):,} ({valuation.get('fsv_pct', 80)}% of OMV)"]
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            # Bold the labels
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_introduction(self, doc: Document, data: Dict[str, Any]):
        """Add introduction and instructions section"""
        doc.add_heading('1. Introduction & Instructions', level=1)
        
        report = data.get('report', {})
        client = data.get('client', {})
        title = data.get('title', {})
        basis_of_value = data.get('basis_of_value', 'Market Value')
        
        intro_text = f"This report is prepared at the request of {client.get('name', 'Client')} "
        intro_text += f"for {report.get('purpose', 'banking and general purposes')}. "
        intro_text += f"The subject property, owned by {title.get('owner_name', 'Property Owner')}, "
        intro_text += f"is valued on the basis of {basis_of_value} in accordance with "
        intro_text += "International Valuation Standards (IVS) and the standards of the "
        intro_text += "Institute of Valuers of Sri Lanka, for the stated purpose only."
        
        doc.add_paragraph(intro_text)
    
    def _add_property_identification(self, doc: Document, data: Dict[str, Any]):
        """Add property identification and title section"""
        doc.add_heading('4. Property Identification & Title', level=1)
        
        id_data = data.get('id', {})
        title = data.get('title', {})
        
        # Identification paragraph
        id_text = f"Identification: The property is identified as Lot {id_data.get('lot_number', 'XX')} "
        id_text += f"in Plan No. {id_data.get('plan_number', 'XXXX')} "
        id_text += f"dated {id_data.get('plan_date', 'XX/XX/XXXX')} "
        id_text += f"prepared by {id_data.get('surveyor_name', 'Licensed Surveyor')}; "
        id_text += f'land known as "{id_data.get('land_name', 'Property Name')}".'
        doc.add_paragraph(id_text)
        
        # Extent
        extent_text = f"Extent: {id_data.get('extent_local', 'XX Perches')} (≈ {id_data.get('extent_metric', 'XX sqm')})."
        doc.add_paragraph(extent_text)
        
        # Boundaries
        boundaries = id_data.get('boundaries', {})
        if boundaries:
            boundary_text = "Boundaries (per plan): "
            boundary_parts = []
            for direction in ['north', 'east', 'south', 'west']:
                if boundaries.get(direction):
                    boundary_parts.append(f"{direction.title()} – {boundaries[direction]}")
            boundary_text += "; ".join(boundary_parts) + "."
            doc.add_paragraph(boundary_text)
        
        # Title information
        title_text = f"Title: {title.get('interest', 'Freehold')} interest held by "
        title_text += f"{title.get('owner_name', 'Owner')} by Deed No. {title.get('deed_no', 'XXXX')} "
        title_text += f"dated {title.get('deed_date', 'XX/XX/XXXX')} ({title.get('notary', 'Notary Public')})."
        doc.add_paragraph(title_text)
    
    def _add_location_description(self, doc: Document, data: Dict[str, Any]):
        """Add situation/location section"""
        doc.add_heading('5. Situation / Location', level=1)
        
        location = data.get('location', {})
        address = data.get('address', {})
        
        location_text = f"The subject property is situated at {address.get('full', 'Property Address')} "
        location_text += f"in {location.get('village', 'Village')}, "
        location_text += f"within the {location.get('gn_division', 'GN Division')} GN Division, "
        location_text += f"{location.get('ds_division', 'DS Division')} Divisional Secretariat, "
        location_text += f"{location.get('district', 'District')} District, "
        location_text += f"{location.get('province', 'Province')} Province."
        
        if location.get('lat') and location.get('lng'):
            location_text += f" Coordinates: {location['lat']}, {location['lng']}."
        
        doc.add_paragraph(location_text)
    
    def _add_access_description(self, doc: Document, data: Dict[str, Any]):
        """Add access section"""
        doc.add_heading('6. Access', level=1)
        
        access = data.get('access', {})
        
        access_text = f"From {access.get('landmark', 'nearby landmark')}, "
        access_text += f"proceed {access.get('directions', 'as per local directions')} "
        access_text += f"via {access.get('road_names', 'local roads')} to reach the property. "
        access_text += f"The final access is {access.get('road_width', 'standard width')} wide "
        access_text += f"{access.get('road_surface', 'metaled')} road with "
        access_text += f"{access.get('maintainer', 'local authority')} maintenance."
        
        doc.add_paragraph(access_text)
    
    def _add_site_description(self, doc: Document, data: Dict[str, Any]):
        """Add site (land) description section"""
        doc.add_heading('7. Site (Land) Description', level=1)
        
        site = data.get('site', {})
        
        site_text = f"The land is {site.get('shape', 'regular')}, "
        site_text += f"{site.get('topography', 'level')} and approximately "
        site_text += f"{site.get('level_vs_road', 'level with')} relative to the access road. "
        site_text += f"Soil is {site.get('soil', 'good bearing capacity')}; "
        site_text += f"water table approx. {site.get('water_table_depth', 'adequate depth')}. "
        site_text += f"Site features include {site.get('features_list', 'standard features')}. "
        
        if site.get('frontage'):
            site_text += f"Frontage approx. {site['frontage']}. "
        
        site_text += f"The site is {site.get('flood_risk', 'low risk')} to flooding/water-logging."
        
        doc.add_paragraph(site_text)
    
    def _add_improvements_description(self, doc: Document, data: Dict[str, Any]):
        """Add improvements/buildings section"""
        buildings = data.get('buildings', [])
        
        if not buildings:
            return
        
        doc.add_heading('8. Improvements / Buildings', level=1)
        
        for building in buildings:
            building_text = f"Building {building.get('index', 1)} – {building.get('type', 'Residential')}: "
            building_text += f"{building.get('storeys', '1')}-storey {building.get('structure', 'masonry')} "
            building_text += f"with {building.get('roof', {}).get('type', 'tiled')} roof on "
            building_text += f"{building.get('roof', {}).get('structure', 'timber')}; "
            building_text += f"walls {building.get('walls', 'masonry')}; "
            building_text += f"floors {building.get('floors', 'cement')}; "
            building_text += f"doors {building.get('doors', 'timber')}; "
            building_text += f"windows {building.get('windows', 'timber')}. "
            building_text += f"Layout: {building.get('layout', 'functional layout')}. "
            building_text += f"Approx. floor area {building.get('area_sqft', 0)} sq.ft "
            building_text += f"(≈ {building.get('area_sqm', 0)} m²). "
            building_text += f"Age {building.get('age', 'moderate')} years; "
            building_text += f"{building.get('condition', 'good')} condition; "
            building_text += f"{building.get('occupancy', 'owner occupied')}."
            
            doc.add_paragraph(building_text)
    
    def _add_services_utilities(self, doc: Document, data: Dict[str, Any]):
        """Add services & utilities section"""
        doc.add_heading('9. Services & Utilities', level=1)
        
        utilities = data.get('utilities', {})
        
        utilities_text = f"Electricity: {utilities.get('electricity', 'Available')}; "
        utilities_text += f"Water: {utilities.get('water', 'Available')}; "
        utilities_text += f"Telephone/Internet: {utilities.get('telecom', 'Available')}; "
        utilities_text += f"Sewerage: {utilities.get('sewerage', 'Standard system')}; "
        utilities_text += f"Drainage: {utilities.get('drainage', 'Adequate')}; "
        utilities_text += f"Other: {utilities.get('other', 'Standard utilities')}."
        
        doc.add_paragraph(utilities_text)
    
    def _add_planning_zoning(self, doc: Document, data: Dict[str, Any]):
        """Add planning/zoning section"""
        doc.add_heading('10. Planning / Zoning / Restrictions', level=1)
        
        planning = data.get('planning', {})
        
        planning_text = f"{planning.get('zoning', 'Residential zoning')}; "
        planning_text += f"Street line/building line: {planning.get('street_line', 'standard setbacks')}; "
        planning_text += f"Reservations / easements: {planning.get('easements', 'No adverse restrictions identified')}."
        
        doc.add_paragraph(planning_text)
    
    def _add_locality_context(self, doc: Document, data: Dict[str, Any]):
        """Add locality & market context section"""
        doc.add_heading('11. Locality & Market Context', level=1)
        
        locality = data.get('locality', {})
        narrative = locality.get('narrative', 'Established area with good amenities and infrastructure, suitable for residential development.')
        
        doc.add_paragraph(narrative)
    
    def _add_valuation_approach(self, doc: Document, data: Dict[str, Any]):
        """Add valuation approach & market evidence section"""
        doc.add_heading('12. Valuation Approach & Market Evidence', level=1)
        
        valuation = data.get('valuation', {})
        evidence = data.get('evidence', {})
        location = data.get('location', {})
        
        approach_text = f"Approach: {valuation.get('approach', 'Comparison method')} "
        approach_text += "(e.g., Comparison method for land; Depreciated Replacement Cost for buildings)."
        doc.add_paragraph(approach_text)
        
        evidence_text = f"Market Evidence: Similar lands ({evidence.get('size_range', 'comparable sizes')}) "
        evidence_text += f"in {location.get('village', 'the area')} / {location.get('district', 'district')} "
        evidence_text += f"transact at Rs. {evidence.get('low', 'market rates')} – {evidence.get('high', 'market rates')} per perch, "
        evidence_text += "adjusted for location/conveniences."
        doc.add_paragraph(evidence_text)
    
    def _add_valuation_calculations(self, doc: Document, data: Dict[str, Any]):
        """Add valuation calculations section"""
        doc.add_heading('13. Valuation Calculations', level=1)
        
        id_data = data.get('id', {})
        calc = data.get('calc', {})
        valuation = data.get('valuation', {})
        buildings = data.get('buildings', [])
        
        # Land calculation
        land_calc = f"Land: {id_data.get('extent_perches', 0)} P @ Rs. {calc.get('land_rate_perch', 0):,} /P = "
        land_calc += f"Rs. {calc.get('land_value', 0):,}"
        doc.add_paragraph(land_calc)
        
        # Building calculations
        for building in buildings:
            building_calc = f"Building {building.get('index', 1)}: {building.get('area_sqft', 0)} sq.ft @ "
            building_calc += f"Rs. {building.get('rate_sqft', 0)} /sq.ft = Rs. {building.get('cost', 0):,}"
            doc.add_paragraph(building_calc)
            
            if building.get('dep_pct', 0) > 0:
                dep_calc = f"Less {building.get('dep_pct', 0)}% depreciation = Rs. {building.get('value_after_dep', 0):,}"
                doc.add_paragraph(dep_calc)
        
        # Other improvements
        if calc.get('other_value', 0) > 0:
            other_calc = f"Other improvements: {calc.get('other_items', 'Various')} = Rs. {calc.get('other_value', 0):,}"
            doc.add_paragraph(other_calc)
        
        # Total values
        doc.add_paragraph()
        total_omv = doc.add_paragraph()
        omv_run = total_omv.add_run(f"Total Open Market Value (OMV) = Rs. {valuation.get('market_value_numeric', 0):,} (rounded)")
        omv_run.font.bold = True
        
        total_fsv = doc.add_paragraph()
        fsv_run = total_fsv.add_run(f"Forced Sale Value (FSV) = Rs. {valuation.get('forced_sale_value_numeric', 0):,} (assumed {valuation.get('fsv_pct', 80)}% of OMV)")
        fsv_run.font.bold = True
        
        # In words
        doc.add_paragraph()
        words_para = doc.add_paragraph()
        words_run = words_para.add_run(f"In Words: {valuation.get('market_value_words', '')}")
        words_run.font.bold = True
    
    def _add_conclusion(self, doc: Document, data: Dict[str, Any]):
        """Add conclusion/opinion of value section"""
        doc.add_heading('14. Conclusion / Opinion of Value', level=1)
        
        valuation = data.get('valuation', {})
        inspection = data.get('inspection', {})
        
        conclusion_text = f"In the valuer's opinion, the Open Market Value of the subject property "
        conclusion_text += f"(land and buildings) as at {inspection.get('date', 'inspection date')} is "
        conclusion_text += f"Rs. {valuation.get('market_value_numeric', 0):,} "
        conclusion_text += f"({valuation.get('market_value_words', '')})."
        
        conclusion_para = doc.add_paragraph()
        conclusion_run = conclusion_para.add_run(conclusion_text)
        conclusion_run.font.bold = True
        
        # Forced sale value if required
        if valuation.get('forced_sale_value_numeric', 0) > 0:
            fsv_text = f"The Forced Sale Value is Rs. {valuation.get('forced_sale_value_numeric', 0):,}."
            doc.add_paragraph(fsv_text)
    
    def _add_certificate_identity(self, doc: Document, data: Dict[str, Any]):
        """Add certificate of identity section"""
        doc.add_heading('15. Certificate of Identity', level=1)
        
        id_data = data.get('id', {})
        
        cert_text = f"I hereby certify that the property inspected and valued herein is "
        cert_text += f"identical to that described as Lot {id_data.get('lot_number', 'XX')} "
        cert_text += f"in Plan No. {id_data.get('plan_number', 'XXXX')} "
        cert_text += f"dated {id_data.get('plan_date', 'XX/XX/XXXX')} "
        cert_text += f"prepared by {id_data.get('surveyor_name', 'Licensed Surveyor')}."
        
        doc.add_paragraph(cert_text)
    
    def _add_assumptions_conditions(self, doc: Document, data: Dict[str, Any]):
        """Add assumptions & limiting conditions section"""
        doc.add_heading('16. Assumptions & Limiting Conditions', level=1)
        
        assumptions = [
            "The valuation assumes a clear, marketable title free of encumbrances unless stated otherwise.",
            "No structural survey or soil investigation was undertaken; buildings are assumed sound commensurate with age and type.",
            f"The valuation is for {data.get('report', {}).get('purpose', 'the stated purpose')} and for the exclusive use of {data.get('client', {}).get('name', 'the client')}; no responsibility to third parties.",
            f"Values are current as at {data.get('inspection', {}).get('date', 'inspection date')} and may change with market conditions.",
            "Measurements are approximate; plans/deeds are relied upon where applicable.",
            "The valuer confirms independence and adherence to IVS/IVSL standards."
        ]
        
        for assumption in assumptions:
            para = doc.add_paragraph()
            para.add_run("• ").font.bold = True
            para.add_run(assumption)
    
    def _add_signature_block(self, doc: Document, data: Dict[str, Any]):
        """Add signature block section"""
        doc.add_heading('17. Signature', level=1)
        
        valuer = data.get('valuer', {})
        report = data.get('report', {})
        
        # Valuer name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(valuer.get('full_name', 'Valuer Name'))
        name_run.font.bold = True
        
        # Qualifications
        quals_para = doc.add_paragraph()
        quals_run = quals_para.add_run(valuer.get('qualifications', 'Professional Qualifications'))
        
        # Registration
        reg_para = doc.add_paragraph()
        reg_run = reg_para.add_run(f"Membership/Registration: {valuer.get('registration_no', 'IVSL Member')}")
        
        # Date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Date: {report.get('date', 'Report Date')}")


# Service instance
def create_template_engine(db: Session) -> ProfessionalTemplateEngine:
    """Factory function to create template engine with database session"""
    return ProfessionalTemplateEngine(db)