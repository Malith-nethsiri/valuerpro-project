"""
Document Generation Service

This module handles the generation of professional property valuation reports
in PDF and DOCX formats using templates and report data.
"""

import os
import json
from datetime import datetime
from typing import Dict, Any, Optional, Union
from pathlib import Path
import tempfile
from io import BytesIO

# Document generation libraries
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from jinja2 import Template

from app.models import Report, User
from app.services.template_engine import create_template_engine
from app.services.number_to_words import convert_lkr_to_currency_words


class DocumentGenerationService:
    """Service for generating professional property valuation documents"""
    
    def __init__(self, db_session=None):
        self.templates_dir = Path(__file__).parent.parent / "templates"
        self.db_session = db_session
        self.ensure_templates_dir()
    
    def ensure_templates_dir(self):
        """Ensure templates directory exists"""
        self.templates_dir.mkdir(exist_ok=True)
    
    def generate_pdf_report(self, report: Report, user: User) -> BytesIO:
        """
        Generate a professional PDF report using ReportLab
        
        Args:
            report: The report model instance
            user: The user who owns the report
            
        Returns:
            BytesIO: PDF content as bytes
        """
        buffer = BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for the 'Flowable' objects
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        # Professional letterhead header
        letterhead_style = ParagraphStyle(
            'LetterheadStyle',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_LEFT,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        contact_style = ParagraphStyle(
            'ContactStyle', 
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_RIGHT,
            textColor=colors.black
        )
        
        # Create valuer letterhead table
        valuer_name = user.full_name
        credentials = []
        if user.title:
            valuer_name = f"{user.title} {valuer_name}"
        if user.qualifications:
            credentials.append(user.qualifications)
        if user.panel_memberships:
            credentials.append(user.panel_memberships)
        
        credential_text = ", ".join(credentials) if credentials else "Chartered Valuer"
        
        # Left side: Valuer info, Right side: Contact info
        contact_info = []
        if user.business_address:
            contact_info.extend(user.business_address.split('\n')[:3])  # Max 3 lines
        if user.contact_numbers:
            contact_info.append(f"Tel: {user.contact_numbers}")
        
        letterhead_data = [[
            f"{valuer_name}\n{credential_text}",
            "\n".join(contact_info) if contact_info else "Contact information available upon request"
        ]]
        
        letterhead_table = Table(letterhead_data, colWidths=[3.5*inch, 2.5*inch])
        letterhead_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (0, 0), 12),
            ('FONTSIZE', (1, 0), (1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
        ]))
        
        story.append(letterhead_table)
        
        # Reference and date line
        current_date = datetime.now()
        ref_number = report.reference_number or f'VRP-{current_date.strftime("%Y%m%d")}-{str(report.id)[:8].upper()}'
        
        ref_date_data = [[
            f"My Ref: {ref_number}",
            f"Date: {current_date.strftime('%d %B %Y')}"
        ]]
        
        ref_table = Table(ref_date_data, colWidths=[3*inch, 3*inch])
        ref_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
        ]))
        
        story.append(ref_table)
        
        # Professional centered title
        story.append(Paragraph("PROPERTY VALUATION REPORT", title_style))
        
        # Subtitle
        subtitle_style = ParagraphStyle(
            'SubtitleStyle',
            parent=styles['Normal'],
            fontSize=14,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("For Banking and General Purposes", subtitle_style))
        story.append(Spacer(1, 30))
        
        # Add report content sections
        if report.data:
            self._add_report_sections_to_pdf(story, report.data, styles, heading_style)
        
        # Professional signature block
        story.append(Spacer(1, 50))
        
        # Signature line and valuer details
        signature_style = ParagraphStyle(
            'SignatureStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_RIGHT,
            spaceAfter=5
        )
        
        # Create signature block
        story.append(Paragraph("_" * 40, signature_style))
        
        # Valuer name and credentials
        valuer_signature = user.full_name
        if user.title:
            valuer_signature = f"{user.title} {valuer_signature}"
        
        story.append(Paragraph(valuer_signature, signature_style))
        
        # Professional credentials
        if user.qualifications:
            story.append(Paragraph(user.qualifications, signature_style))
        else:
            story.append(Paragraph("Chartered Valuer", signature_style))
        
        # Institute membership
        if user.panel_memberships:
            story.append(Paragraph(user.panel_memberships, signature_style))
        
        story.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _add_report_sections_to_pdf(self, story: list, data: Dict[str, Any], styles, heading_style):
        """Add professional report content sections following Sri Lankan valuation standards"""
        
        # Safe data access helpers - handle None data
        if not data or not isinstance(data, dict):
            data = {}
            
        property_details = data.get('property_details') or {}
        location_data = data.get('location_data') or {}
        property_info = data.get('property_info') or {}
        valuation = data.get('valuation') or {}
        ai_analysis = data.get('ai_analysis') or {}
        
        # Helper function to safely handle Unicode strings for ReportLab
        def safe_str(value, default=''):
            if value is None or value == 'None':
                return default
            try:
                text = str(value).strip()
                # Replace problematic Unicode characters with ASCII equivalents
                text = text.encode('ascii', 'replace').decode('ascii')
                return text if text else default
            except Exception:
                return default
        
        # Helper to format extent in both local and metric units
        def format_extent(extent_str):
            if not extent_str or extent_str == 'None':
                return 'Not specified'
            try:
                # If it's already formatted with perches, return as is
                if 'perch' in extent_str.lower():
                    return extent_str
                # Try to parse as numeric value in perches
                perches = float(extent_str)
                acres = int(perches // 160)
                roods = int((perches % 160) // 40)
                remaining_perches = perches % 40
                
                local_format = f"{acres}A-{roods}R-{remaining_perches:.1f}P"
                hectares = perches * 0.00253
                return f"{local_format} (approximately {hectares:.4f} hectares)"
            except:
                return extent_str
        
        # Professional Executive Summary
        story.append(Paragraph("EXECUTIVE SUMMARY", heading_style))
        
        # Get property type from data
        prop_type = safe_str(property_details.get('property_type'), 'Residential Property')
        if ai_analysis.get('building_details', {}).get('building_type'):
            building_type = safe_str(ai_analysis['building_details']['building_type'])
            if building_type and building_type.lower() != 'none':
                prop_type = f"Land with {building_type}"
        
        # Format location from location data
        location_parts = []
        if location_data.get('components'):
            components = location_data['components']
            if components.get('village'):
                location_parts.append(safe_str(components['village']))
            if components.get('district'):
                location_parts.append(f"{safe_str(components['district'])} District")
        
        location_text = ', '.join(location_parts) if location_parts else safe_str(location_data.get('formatted_address'), 'Location not specified')
        
        # Create professional summary table
        summary_data = [
            ['Property Type:', prop_type],
            ['Location:', location_text],
            ['Land Extent:', format_extent(safe_str(property_info.get('extent')))],
            ['Market Value:', f"Rs. {valuation.get('total_market_value', 0):,}.00" if valuation.get('total_market_value') else 'To be assessed'],
            ['Forced Sale Value:', f"Rs. {int(valuation.get('total_market_value', 0) * 0.8):,}.00" if valuation.get('total_market_value') else 'To be assessed (80% of Market Value)'],
        ]
        
        summary_table = Table(summary_data, colWidths=[2.2*inch, 3.8*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 30))
        
        # 1. INTRODUCTION
        story.append(Paragraph("1. INTRODUCTION", heading_style))
        
        # Professional introduction paragraph
        intro_text = (
            "This valuation report is prepared in accordance with International Valuation Standards (IVS) "
            "and the professional standards of the Institute of Valuers of Sri Lanka. The purpose of this "
            "valuation is for banking and general purposes as requested by the client. The interest valued "
            "is the freehold interest in the subject property, and the basis of valuation is Market Value."
        )
        story.append(Paragraph(intro_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # 2. PROPERTY IDENTIFICATION
        story.append(Paragraph("2. PROPERTY IDENTIFICATION", heading_style))
        
        # Professional narrative for property identification
        identification_parts = []
        
        if property_info.get('lot_number') and property_info.get('plan_number'):
            lot_num = safe_str(property_info['lot_number'])
            plan_num = safe_str(property_info['plan_number'])
            plan_date = safe_str(property_info.get('plan_date'))
            surveyor = safe_str(property_info.get('surveyor_name'))
            
            id_text = f"The subject property is identified as Lot {lot_num} in Plan No. {plan_num}"
            if plan_date:
                id_text += f" dated {plan_date}"
            if surveyor:
                id_text += f", surveyed by {surveyor}, Licensed Surveyor"
            id_text += "."
            identification_parts.append(id_text)
        
        # Land name
        if property_info.get('land_name'):
            land_name = safe_str(property_info['land_name'])
            if land_name and land_name.lower() != 'none':
                identification_parts.append(f'The land is named "{land_name}" as per the survey documents.')
        
        # Extent
        if property_info.get('extent'):
            extent_text = f"The extent of the land is {format_extent(safe_str(property_info['extent']))}."
            identification_parts.append(extent_text)
        
        for part in identification_parts:
            story.append(Paragraph(part, styles['Normal']))
        
        story.append(Spacer(1, 15))
        
        # Boundaries in narrative format
        boundaries = property_info.get('boundaries') or {}
        if boundaries and any(boundaries.values()):
            story.append(Paragraph("Boundaries:", styles['Heading3']))
            boundary_parts = []
            for direction in ['north', 'south', 'east', 'west']:
                boundary = safe_str(boundaries.get(direction))
                if boundary and boundary.lower() != 'n/a':
                    boundary_parts.append(f"{direction.capitalize()}: {boundary}")
            
            if boundary_parts:
                boundary_text = "; ".join(boundary_parts) + "."
                story.append(Paragraph(boundary_text, styles['Normal']))
        
        story.append(Spacer(1, 20))
        
        # 3. SITUATION AND LOCATION
        story.append(Paragraph("3. SITUATION AND LOCATION", heading_style))
        
        # Professional location description
        location_parts = []
        
        if location_data.get('components'):
            components = location_data['components']
            location_desc = "The subject property is situated"
            
            if components.get('village'):
                village = safe_str(components['village'])
                location_desc += f" in {village}"
            
            if components.get('district'):
                district = safe_str(components['district'])
                location_desc += f" within {district} District"
            
            if components.get('province'):
                province = safe_str(components['province'])
                location_desc += f", {province} Province"
            
            location_desc += ", Democratic Socialist Republic of Sri Lanka."
            story.append(Paragraph(location_desc, styles['Normal']))
        
        # Access route if available
        if location_data.get('access_route', {}).get('description'):
            access_desc = safe_str(location_data['access_route']['description'])
            if access_desc:
                story.append(Paragraph(f"Access: {access_desc}", styles['Normal']))
        
        story.append(Spacer(1, 20))
        
        # 4. DESCRIPTION OF PROPERTY
        story.append(Paragraph("4. DESCRIPTION OF PROPERTY", heading_style))
        
        # Land description
        story.append(Paragraph("4.1 Description of Land:", styles['Heading3']))
        
        land_desc_parts = []
        
        # Basic land characteristics
        if property_info.get('extent'):
            extent_str = format_extent(safe_str(property_info['extent']))
            land_desc_parts.append(f"The land comprises an extent of {extent_str}.")
        
        # Additional features from AI analysis
        if ai_analysis.get('additional_features'):
            features = ai_analysis['additional_features']
            if isinstance(features, list) and features:
                feature_text = "The property features include: "
                clean_features = [f for f in features if f and str(f).lower() not in ['none', 'null']]
                if clean_features:
                    feature_text += ", ".join(clean_features).lower() + "."
                    land_desc_parts.append(feature_text)
        
        # Default land description if no specific details
        if not land_desc_parts:
            land_desc_parts.append("The land is a regular block of land suitable for residential development.")
        
        for part in land_desc_parts:
            story.append(Paragraph(part, styles['Normal']))
        
        story.append(Spacer(1, 15))
        
        # Building description if available
        if ai_analysis.get('building_details'):
            building_details = ai_analysis['building_details']
            if building_details.get('building_type') and str(building_details['building_type']).lower() not in ['none', 'null']:
                story.append(Paragraph("4.2 Description of Buildings:", styles['Heading3']))
                
                building_type = safe_str(building_details['building_type'])
                building_text = f"The property contains a {building_type.lower()}"
                
                if building_details.get('floors') and str(building_details['floors']).lower() != 'none':
                    floors = safe_str(building_details['floors'])
                    building_text += f" of {floors} storey(s)"
                
                building_text += ". The building is constructed of conventional materials and is in serviceable condition."
                
                story.append(Paragraph(building_text, styles['Normal']))
        
        story.append(Spacer(1, 20))
        
        # 5. VALUATION APPROACH AND CALCULATIONS
        story.append(Paragraph("5. VALUATION APPROACH AND CALCULATIONS", heading_style))
        
        # Professional valuation approach description
        approach_text = (
            "The valuation is carried out using the Comparison Method, being the most appropriate "
            "method for residential properties in this locality. The Open Market Value is assessed "
            "considering recent comparable sales and current market conditions in the area."
        )
        story.append(Paragraph(approach_text, styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Valuation calculations
        if valuation and valuation.get('total_market_value'):
            story.append(Paragraph("Valuation Calculations:", styles['Heading3']))
            
            market_value = valuation['total_market_value']
            
            # Simple calculation display (can be enhanced with detailed breakdown)
            calc_data = [
                ['Land Value:', f"Rs. {int(market_value * 0.6):,}.00"],
                ['Building Value:', f"Rs. {int(market_value * 0.4):,}.00"],
                ['Total Market Value:', f"Rs. {market_value:,}.00"],
                ['Forced Sale Value (80%):', f"Rs. {int(market_value * 0.8):,}.00"],
            ]
            
            calc_table = Table(calc_data, colWidths=[2.5*inch, 2.5*inch])
            calc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (0, 2), (1, 2), 'Helvetica-Bold'),  # Total row bold
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('LINEABOVE', (0, 2), (-1, 2), 2, colors.black),  # Line above total
            ]))
            story.append(calc_table)
        else:
            # Placeholder for manual valuation input
            placeholder_text = (
                "Valuation calculations to be inserted based on detailed analysis of land rates, "
                "building costs, and market conditions in the locality."
            )
            story.append(Paragraph(placeholder_text, styles['Italic']))
        
        story.append(Spacer(1, 30))
        
        # 6. CONCLUSION
        story.append(Paragraph("6. CONCLUSION", heading_style))
        
        # Professional conclusion
        if valuation and valuation.get('total_market_value'):
            market_value = valuation['total_market_value']
            
            # Convert number to words (simplified version)
            def number_to_words_simple(amount):
                if amount >= 1000000:
                    millions = amount // 1000000
                    remainder = amount % 1000000
                    if remainder >= 100000:
                        hundreds_k = remainder // 100000
                        return f"{millions} Million {hundreds_k} Hundred Thousand"
                    else:
                        return f"{millions} Million"
                elif amount >= 100000:
                    hundreds_k = amount // 100000
                    return f"{hundreds_k} Hundred Thousand"
                else:
                    return str(amount)
            
            value_words = number_to_words_simple(int(market_value))
            
            conclusion_text = (
                f"Based on the foregoing analysis and considering the prevailing market conditions, "
                f"I am of the opinion that the Open Market Value of the subject property as at the "
                f"date of inspection is Rs. {market_value:,}.00 (Sri Lanka Rupees {value_words} only)."
            )
            story.append(Paragraph(conclusion_text, styles['Normal']))
            
            # Forced sale value conclusion
            forced_sale_value = int(market_value * 0.8)
            forced_sale_words = number_to_words_simple(forced_sale_value)
            
            story.append(Spacer(1, 10))
            forced_sale_text = (
                f"For forced sale purposes, assuming a disposal period of 3-6 months, "
                f"the value is assessed at Rs. {forced_sale_value:,}.00 "
                f"(Sri Lanka Rupees {forced_sale_words} only)."
            )
            story.append(Paragraph(forced_sale_text, styles['Normal']))
        else:
            conclusion_text = (
                "The Open Market Value of the subject property will be determined upon completion "
                "of the detailed valuation analysis and market research."
            )
            story.append(Paragraph(conclusion_text, styles['Normal']))
        
        story.append(Spacer(1, 30))
        
        # 7. CERTIFICATE OF IDENTIFICATION
        story.append(Paragraph("7. CERTIFICATE OF IDENTIFICATION", heading_style))
        
        # Professional certificate
        if property_info.get('lot_number') and property_info.get('plan_number'):
            lot_num = safe_str(property_info['lot_number'])
            plan_num = safe_str(property_info['plan_number'])
            plan_date = safe_str(property_info.get('plan_date', ''))
            surveyor = safe_str(property_info.get('surveyor_name', ''))
            
            cert_text = (
                f"I hereby certify that the property inspected by me and valued herein is "
                f"identical to the land depicted as Lot {lot_num} in Plan No. {plan_num}"
            )
            if plan_date:
                cert_text += f" dated {plan_date}"
            if surveyor:
                cert_text += f" by {surveyor}, Licensed Surveyor"
            cert_text += "."
        else:
            cert_text = (
                "I hereby certify that the property inspected by me and valued herein "
                "corresponds to the property described in the available title documents."
            )
        
        story.append(Paragraph(cert_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # 8. ASSUMPTIONS AND LIMITING CONDITIONS
        story.append(Paragraph("8. ASSUMPTIONS AND LIMITING CONDITIONS", heading_style))
        
        assumptions = [
            "This valuation is carried out on the assumption of a clear marketable title free from encumbrances.",
            "No responsibility is assumed for matters legal in nature affecting the title or for hidden defects in the property.",
            "The valuation is valid as of the date of inspection and may be affected by subsequent market changes.",
            "This report is prepared for the stated purpose and should not be used for any other purpose without the valuer's consent.",
            "No detailed structural survey was conducted; the valuation assumes the buildings are in reasonable condition for their age.",
        ]
        
        for assumption in assumptions:
            story.append(Paragraph(f"• {assumption}", styles['Normal']))
        
        story.append(Spacer(1, 30))
    
    def generate_docx_report(self, report: Report, user: User) -> BytesIO:
        """
        Generate a professional DOCX report using python-docx with narrative structure
        
        Args:
            report: The report model instance
            user: The user who owns the report
            
        Returns:
            BytesIO: DOCX content as bytes
        """
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        # Create document
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Helper functions
        def safe_str(value, default='N/A'):
            if value is None:
                return default
            if isinstance(value, str) and value.strip():
                return value.strip()
            elif value:
                return str(value)
            return default
        
        def add_heading_custom(text, level=1):
            heading = doc.add_heading(text, level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            return heading
        
        def add_paragraph_custom(text, bold=False, italic=False):
            para = doc.add_paragraph()
            run = para.add_run(text)
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            return para
        
        # Extract data safely
        data = report.data or {}
        property_info = data.get('property_info', {})
        location_data = data.get('location_data', {})
        property_details = data.get('property_details', {})
        valuation = data.get('valuation', {})
        
        # === PROFESSIONAL LETTERHEAD ===
        # Main title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("PROPERTY VALUATION REPORT")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Valuer details
        doc.add_paragraph()
        valuer_para = doc.add_paragraph()
        valuer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        valuer_text = f"Prepared by: {user.title or ''} {user.full_name}\n"
        if user.qualifications:
            valuer_text += f"{user.qualifications}\n"
        if user.panel_memberships:
            valuer_text += f"{user.panel_memberships}\n"
        if user.business_address:
            valuer_text += f"{user.business_address}\n"
        if user.contact_numbers:
            valuer_text += f"Contact: {user.contact_numbers}\n"
        
        valuer_run = valuer_para.add_run(valuer_text.strip())
        valuer_run.font.name = 'Arial'
        valuer_run.font.size = Pt(11)
        
        # Report reference and date
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_text = f"Report Reference: {report.reference_number or str(report.id)[:8]}\n"
        ref_text += f"Date: {report.created_at.strftime('%B %d, %Y')}"
        ref_run = ref_para.add_run(ref_text)
        ref_run.font.name = 'Arial'
        ref_run.font.size = Pt(10)
        
        doc.add_page_break()
        
        # === MAIN REPORT CONTENT ===
        
        # 1. INTRODUCTION
        add_heading_custom("1. INTRODUCTION", 1)
        intro_text = f"This valuation report has been prepared for {safe_str(report.title)} "
        intro_text += f"regarding the property located at {safe_str(report.property_address)}. "
        intro_text += f"The inspection was carried out and this report was prepared on {report.created_at.strftime('%B %d, %Y')}."
        add_paragraph_custom(intro_text)
        
        # 2. PROPERTY IDENTIFICATION
        add_heading_custom("2. PROPERTY IDENTIFICATION", 1)
        
        # Extract key identifiers
        lot_num = safe_str(property_info.get('lot_number'))
        plan_num = safe_str(property_info.get('plan_number'))
        plan_date = safe_str(property_info.get('plan_date'))
        surveyor = safe_str(property_info.get('surveyor'))
        
        if lot_num != 'N/A' and plan_num != 'N/A':
            id_text = f"The subject property is identified as Lot {lot_num} in Plan No. {plan_num}"
            if plan_date != 'N/A':
                id_text += f" dated {plan_date}"
            if surveyor != 'N/A':
                id_text += f", prepared by Licensed Surveyor {surveyor}"
            id_text += "."
            add_paragraph_custom(id_text)
        
        # Property extent
        extent = safe_str(property_info.get('extent'))
        if extent != 'N/A':
            add_paragraph_custom(f"The total extent of the property is {extent}.")
        
        # Boundaries
        boundaries = property_info.get('boundaries', {})
        if any(boundaries.values()):
            add_paragraph_custom("The property boundaries are described as follows:", bold=True)
            for direction in ['north', 'south', 'east', 'west']:
                boundary = safe_str(boundaries.get(direction))
                if boundary != 'N/A':
                    add_paragraph_custom(f"• {direction.title()}: {boundary}")
        
        # 3. LOCATION AND ACCESS
        add_heading_custom("3. LOCATION AND ACCESS", 1)
        
        formatted_address = safe_str(location_data.get('formatted_address'))
        if formatted_address != 'N/A':
            add_paragraph_custom(f"The property is situated at {formatted_address}.")
        
        # Add route description if available
        route = safe_str(location_data.get('access_route'))
        if route != 'N/A':
            add_paragraph_custom(f"Access Route: {route}")
        
        # Administrative area information
        district = safe_str(location_data.get('administrative_area_level_2'))
        province = safe_str(location_data.get('administrative_area_level_1'))
        if district != 'N/A' or province != 'N/A':
            admin_text = "The property is administratively located in "
            if district != 'N/A':
                admin_text += f"{district} District"
            if province != 'N/A':
                if district != 'N/A':
                    admin_text += f", {province} Province"
                else:
                    admin_text += f"{province} Province"
            admin_text += "."
            add_paragraph_custom(admin_text)
        
        # 4. PROPERTY DESCRIPTION
        add_heading_custom("4. PROPERTY DESCRIPTION", 1)
        
        # Property type and nature
        prop_type = safe_str(property_details.get('property_type'), 'Residential Land')
        add_paragraph_custom(f"The subject property is classified as {prop_type}.")
        
        # Additional property details
        for key, value in property_details.items():
            if key != 'property_type' and value:
                detail_text = f"{key.replace('_', ' ').title()}: {safe_str(value)}"
                add_paragraph_custom(detail_text)
        
        # 5. MARKET ANALYSIS AND VALUATION
        add_heading_custom("5. MARKET ANALYSIS AND VALUATION", 1)
        
        # Valuation methodology
        add_paragraph_custom("This valuation has been prepared using appropriate valuation methodologies including market comparison approach and considering current market conditions.")
        
        # Market value
        market_value = valuation.get('market_value')
        if market_value:
            if isinstance(market_value, (int, float)):
                value_text = f"Based on our analysis, the current market value of the property is estimated at Rs. {market_value:,}."
            else:
                value_text = f"Market Value: {safe_str(market_value)}"
            add_paragraph_custom(value_text, bold=True)
        
        # Other valuation details
        for key, value in valuation.items():
            if key != 'market_value' and value:
                val_text = f"{key.replace('_', ' ').title()}: "
                if 'value' in key.lower() or 'price' in key.lower():
                    if isinstance(value, (int, float)):
                        val_text += f"Rs. {value:,}"
                    else:
                        val_text += safe_str(value)
                else:
                    val_text += safe_str(value)
                add_paragraph_custom(val_text)
        
        # 6. ASSUMPTIONS AND CONDITIONS
        add_heading_custom("6. ASSUMPTIONS AND CONDITIONS", 1)
        add_paragraph_custom("This valuation is subject to the following assumptions and conditions:")
        add_paragraph_custom("• The property has clear and marketable title")
        add_paragraph_custom("• All statutory approvals are in place")
        add_paragraph_custom("• The property is free from any encumbrances")
        add_paragraph_custom("• Physical inspection limitations, if any, have been noted")
        
        # 7. CONCLUSION
        add_heading_custom("7. CONCLUSION", 1)
        conclusion_text = f"Based on our professional analysis and market research, the fair market value "
        conclusion_text += f"of the property located at {safe_str(report.property_address)} "
        if market_value:
            if isinstance(market_value, (int, float)):
                conclusion_text += f"is Rs. {market_value:,} as of {report.created_at.strftime('%B %d, %Y')}."
            else:
                conclusion_text += f"has been assessed and detailed above as of {report.created_at.strftime('%B %d, %Y')}."
        else:
            conclusion_text += f"has been professionally assessed as detailed in this report as of {report.created_at.strftime('%B %d, %Y')}."
        add_paragraph_custom(conclusion_text)
        
        # 8. CERTIFICATION
        add_heading_custom("8. VALUER CERTIFICATION", 1)
        cert_text = f"I, {user.title or ''} {user.full_name}"
        if user.qualifications:
            cert_text += f", {user.qualifications}"
        cert_text += ", hereby certify that:"
        add_paragraph_custom(cert_text)
        
        certifications = [
            "This valuation has been prepared in accordance with professional standards",
            "The valuation reflects my professional opinion of market value", 
            "I have no personal interest in the subject property",
            "This report is prepared for the stated purpose only"
        ]
        
        for cert in certifications:
            add_paragraph_custom(f"• {cert}")
        
        # Signature block
        doc.add_paragraph()
        doc.add_paragraph()
        sig_para = add_paragraph_custom("_" * 30)
        add_paragraph_custom(f"{user.title or ''} {user.full_name}")
        if user.qualifications:
            add_paragraph_custom(user.qualifications)
        add_paragraph_custom(f"Date: {report.created_at.strftime('%B %d, %Y')}")
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_professional_docx_report(self, report_id: str) -> BytesIO:
        """
        Generate a professional DOCX report using the new template engine
        with complete database mapping and professional formatting
        
        Args:
            report_id: The report ID to generate
            
        Returns:
            BytesIO: DOCX content as bytes
        """
        if not self.db_session:
            raise ValueError("Database session required for professional DOCX generation")
        
        # Use the new template engine
        template_engine = create_template_engine(self.db_session)
        return template_engine.generate_professional_docx(report_id)
    
    def generate_enhanced_pdf_report(self, report_id: str) -> BytesIO:
        """
        Generate an enhanced PDF report with improved number formatting
        and professional structure
        
        Args:
            report_id: The report ID to generate
            
        Returns:
            BytesIO: PDF content as bytes
        """
        if not self.db_session:
            # Fallback to old method if no DB session
            report = Report(id=report_id, title="Sample Report", property_address="Sample Address")
            user = User(full_name="Sample Valuer")
            return self.generate_pdf_report(report, user)
        
        # Get report data using template engine mapper
        template_engine = create_template_engine(self.db_session)
        template_data = template_engine.mapper.map_report_data(report_id)
        
        # Generate PDF with mapped data
        return self._generate_pdf_with_template_data(template_data)
    
    def _generate_pdf_with_template_data(self, template_data: Dict[str, Any]) -> BytesIO:
        """
        Generate PDF using template data with enhanced formatting
        """
        buffer = BytesIO()
        
        # Create the PDF document with professional styling
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Enhanced title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'], 
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        # Professional letterhead using template data
        valuer = template_data.get('valuer', {})
        report = template_data.get('report', {})
        
        # Letterhead table
        letterhead_data = [[
            f"{valuer.get('titles', '')} {valuer.get('full_name', 'Valuer Name')}\n{valuer.get('qualifications', 'Professional Qualifications')}",
            f"{valuer.get('address_multiline', 'Address')}\nTel: {valuer.get('phones_list', 'Phone')}\nEmail: {valuer.get('email', 'Email')}"
        ]]
        
        letterhead_table = Table(letterhead_data, colWidths=[3.5*inch, 2.5*inch])
        letterhead_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (0, 0), 12),
            ('FONTSIZE', (1, 0), (1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
        ]))
        
        story.append(letterhead_table)
        
        # Reference and date
        ref_data = [[
            f"My Ref: {report.get('ref', 'REF-001')}",
            f"Date: {report.get('date', datetime.now().strftime('%d %B %Y'))}"
        ]]
        
        ref_table = Table(ref_data, colWidths=[3*inch, 3*inch])
        ref_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
        ]))
        
        story.append(ref_table)
        
        # Title
        story.append(Paragraph("VALUATION REPORT", title_style))
        
        # Subject line using template data
        id_data = template_data.get('id', {})
        subject_text = f"Valuation Report of the Property depicted as Lot {id_data.get('lot_number', 'XX')} "
        subject_text += f"in Plan No. {id_data.get('plan_number', 'XXXX')} dated {id_data.get('plan_date', 'XX/XX/XXXX')}"
        
        story.append(Paragraph(subject_text, styles['Normal']))
        story.append(Spacer(1, 30))
        
        # Executive Summary
        story.append(Paragraph("EXECUTIVE SUMMARY", heading_style))
        
        summary = template_data.get('summary', {})
        location = template_data.get('location', {})
        valuation = template_data.get('valuation', {})
        
        summary_data = [
            ['Property Type:', summary.get('property_type', 'Residential Property')],
            ['Location:', f"{location.get('village', '')}, {location.get('district', '')} District, {location.get('province', '')}"],
            ['Land Extent:', f"{id_data.get('extent_local', 'XX Perches')} (≈ {id_data.get('extent_metric', 'XX sqm')})"],
            ['Open Market Value (OMV):', f"Rs. {valuation.get('market_value_numeric', 0):,}"],
            ['Forced Sale Value (FSV):', f"Rs. {valuation.get('forced_sale_value_numeric', 0):,} ({valuation.get('fsv_pct', 80)}% of OMV)"]
        ]
        
        summary_table = Table(summary_data, colWidths=[2.2*inch, 3.8*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 30))
        
        # Main sections with template data
        self._add_enhanced_pdf_sections(story, template_data, styles, heading_style)
        
        # Professional signature with template data
        story.append(Spacer(1, 50))
        
        signature_style = ParagraphStyle(
            'SignatureStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_RIGHT,
            spaceAfter=5
        )
        
        story.append(Paragraph("_" * 40, signature_style))
        story.append(Paragraph(f"{valuer.get('titles', '')} {valuer.get('full_name', 'Valuer Name')}", signature_style))
        story.append(Paragraph(valuer.get('qualifications', 'Professional Qualifications'), signature_style))
        story.append(Paragraph(f"Registration: {valuer.get('registration_no', 'IVSL Member')}", signature_style))
        story.append(Paragraph(f"Date: {report.get('date', datetime.now().strftime('%d %B %Y'))}", signature_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _add_enhanced_pdf_sections(self, story: list, template_data: Dict[str, Any], styles, heading_style):
        """Add enhanced PDF sections using template data"""
        
        # Introduction
        story.append(Paragraph("1. INTRODUCTION", heading_style))
        
        report = template_data.get('report', {})
        client = template_data.get('client', {})
        title = template_data.get('title', {})
        
        intro_text = f"This report is prepared at the request of {client.get('name', 'Client')} "
        intro_text += f"for {report.get('purpose', 'banking and general purposes')}. "
        intro_text += f"The subject property, owned by {title.get('owner_name', 'Property Owner')}, "
        intro_text += f"is valued on the basis of {template_data.get('basis_of_value', 'Market Value')} "
        intro_text += "in accordance with International Valuation Standards (IVS) and the standards of the "
        intro_text += "Institute of Valuers of Sri Lanka, for the stated purpose only."
        
        story.append(Paragraph(intro_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Property Identification
        story.append(Paragraph("2. PROPERTY IDENTIFICATION", heading_style))
        
        id_data = template_data.get('id', {})
        id_text = f"The property is identified as Lot {id_data.get('lot_number', 'XX')} "
        id_text += f"in Plan No. {id_data.get('plan_number', 'XXXX')} "
        id_text += f"dated {id_data.get('plan_date', 'XX/XX/XXXX')} "
        id_text += f"prepared by {id_data.get('surveyor_name', 'Licensed Surveyor')}."
        
        story.append(Paragraph(id_text, styles['Normal']))
        
        extent_text = f"Extent: {id_data.get('extent_local', 'XX Perches')} (≈ {id_data.get('extent_metric', 'XX sqm')})."
        story.append(Paragraph(extent_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Valuation Conclusion with proper number formatting
        story.append(Paragraph("3. VALUATION CONCLUSION", heading_style))
        
        valuation = template_data.get('valuation', {})
        market_value = valuation.get('market_value_numeric', 0)
        
        conclusion_text = f"Based on our professional analysis, the Open Market Value of the subject property "
        conclusion_text += f"as at {template_data.get('inspection', {}).get('date', 'inspection date')} is "
        conclusion_text += f"Rs. {market_value:,} "
        conclusion_text += f"({valuation.get('market_value_words', convert_lkr_to_currency_words(market_value))})."
        
        story.append(Paragraph(conclusion_text, styles['Normal']))
        
        if valuation.get('forced_sale_value_numeric', 0) > 0:
            fsv_value = valuation.get('forced_sale_value_numeric', 0)
            fsv_text = f"The Forced Sale Value is Rs. {fsv_value:,} "
            fsv_text += f"({convert_lkr_to_currency_words(fsv_value)})."
            story.append(Paragraph(fsv_text, styles['Normal']))
        
        story.append(Spacer(1, 20))
        
        # Certificate of Identity
        story.append(Paragraph("4. CERTIFICATE OF IDENTITY", heading_style))
        
        cert_text = f"I hereby certify that the property inspected and valued herein is "
        cert_text += f"identical to that described as Lot {id_data.get('lot_number', 'XX')} "
        cert_text += f"in Plan No. {id_data.get('plan_number', 'XXXX')} "
        cert_text += f"dated {id_data.get('plan_date', 'XX/XX/XXXX')} "
        cert_text += f"prepared by {id_data.get('surveyor_name', 'Licensed Surveyor')}."
        
        story.append(Paragraph(cert_text, styles['Normal']))


# Service instance
document_service = DocumentGenerationService()